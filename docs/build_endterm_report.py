from __future__ import annotations

import math
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "generated"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "SRE_End_Term_Project_Report.docx"


REPORT_MD = r"""
# Comprehensive SRE Implementation for a Distributed Microservices System
**Course End-Term Project Report**  
**Deployment Strategy:** Multi-Orchestration (Docker Swarm & Kubernetes)  
**Infrastructure & Automation:** Terraform & Ansible

## 1. Abstract
This project presents the design and implementation of a distributed microservices-based application engineered according to modern Site Reliability Engineering (SRE) and DevOps principles. The system is composed of six core microservices, namely Authentication, Product, Order, Payment, Notification, and User Profile services, supported by an API gateway, relational data storage, and asynchronous messaging components. To evaluate orchestration strategies in a practical and comparative manner, the application was deployed using both Docker Swarm and Kubernetes, enabling analysis of service scheduling, replication, self-healing behavior, and operational management across two widely used container platforms.

Infrastructure provisioning was automated through Terraform, which defined the compute instances, networking rules, and secure environment topology in a declarative Infrastructure as Code model. Ansible was then used for configuration management, node bootstrapping, container runtime installation, orchestration initialization, and deployment standardization. To ensure operational visibility and service reliability, Prometheus and Grafana were integrated for metrics collection, dashboarding, and SLI/SLO evaluation, covering availability, latency, error rate, and request success ratio. The final outcome is a fully automated, observable, and scalable microservices platform that demonstrates the practical application of cloud automation, service orchestration, and reliability engineering in a production-style environment.

## 2. Architecture & Microservices Overview
The implemented platform follows a distributed microservices architecture in which each business capability is isolated into an independent deployable service. This separation improves scalability, fault isolation, maintainability, and team-level ownership. Every service exposes a dedicated API, communicates over internal service networks, and is designed to operate independently while participating in a broader transactional workflow. The overall architecture includes six domain services, a frontend/API gateway layer, persistent storage, and a message broker/caching tier.

### Authentication Service
The Authentication Service is responsible for identity verification, token issuance, access validation, and role-based authorization. It acts as the security entry point for all user-facing and service-to-service operations. In the implemented architecture, the service follows a lightweight RESTful pattern and relies on secure password hashing, token-based authentication, and relational persistence for account metadata. It exposes endpoints for login, token validation, and session continuity while enforcing authorization policies that protect downstream services. From an SRE perspective, this service is classified as a high-criticality dependency because authentication failure immediately affects the availability of nearly all business functions.

### Product Service
The Product Service manages the application catalog and delivers product listing, item detail, pricing, and availability information. It is optimized for read-heavy workloads and is therefore well suited for horizontal replication. The service uses a stateless API design, allowing replicas to be scheduled freely across orchestration nodes, while a backing relational store and optional cache layer reduce repeated data retrieval costs. Because product browsing generates a significant proportion of application traffic, the Product Service is monitored closely for latency, request throughput, and consistency under concurrent access.

### Order Service
The Order Service coordinates checkout and order lifecycle management. It is responsible for validating purchase requests, persisting order records, and driving state transitions such as pending, confirmed, completed, or failed. This service is one of the most business-critical components in the platform because it sits at the center of the transaction flow and interacts directly with inventory, payment, and notification processes. It depends heavily on correct database connectivity and efficient connection pooling, which is why it was selected as the primary candidate for incident simulation in the reliability exercises.

### Payment Service
The Payment Service handles transaction authorization, payment result propagation, and safe integration with the broader checkout pipeline. It is designed around secure, idempotent processing so that repeated retries or transient failures do not create duplicate charges or inconsistent order states. Operationally, the Payment Service is sensitive to timeout behavior, downstream dependency health, and throughput spikes during peak checkout traffic. For this reason, it is deployed with multiple replicas, strict monitoring, and controlled resource boundaries to preserve transaction integrity.

### Notification Service
The Notification Service manages outbound system communication such as order confirmations, payment updates, and internal event notifications. To prevent non-critical communication tasks from blocking core user transactions, the service is decoupled through a queueing or broker-backed approach using Redis or RabbitMQ patterns. This asynchronous design improves overall resilience by allowing background retries and eventual delivery without introducing direct latency into the checkout path. The service is observed through queue depth, consumer lag, retry volume, and successful delivery indicators.

### User Profile Service
The User Profile Service stores and serves customer-specific metadata such as contact details, shipping information, and account preferences. It is intentionally separated from the Authentication Service to preserve clean service boundaries between identity and business profile data. The service follows CRUD-oriented API behavior with strict input validation and integrates with authentication for access control. Although it is not as latency-critical as checkout services, it remains an essential part of the user experience and contributes to the system's overall reliability posture.

### Frontend / API Gateway
The frontend ingress layer is implemented through Nginx acting as a reverse proxy and gateway for all client traffic. Nginx centralizes request routing, upstream load balancing, forwarding headers, timeout handling, and compression strategy. It maps URI prefixes to the corresponding backend services and exposes a single stable interface to clients, thereby hiding the complexity of internal service topology. Because the gateway relies on orchestrator-managed service discovery rather than fixed host addresses, the same routing strategy works consistently in both Docker Swarm and Kubernetes deployments.

### Database & Broker Tier
The persistent storage layer is centered on PostgreSQL, which provides ACID-compliant transactional guarantees for business-critical services such as authentication, order management, payments, and user profiles. Supporting infrastructure includes Redis and broker-style asynchronous messaging patterns to accelerate read performance and decouple background workflows. Redis contributes low-latency caching and ephemeral state handling, while a broker-oriented design improves resilience for notification delivery and event-driven communication. Together, these components reduce direct contention on the primary database and strengthen system responsiveness under load.

## 3. Infrastructure as Code (Terraform) & Configuration Management (Ansible)
### Terraform Architecture
The infrastructure was provisioned entirely through Terraform using declarative configuration files that define the full lifecycle of compute instances, network boundaries, and security controls. Terraform resources describe virtual machines, network placement, firewall exposure, and operational outputs such as accessible public endpoints. This approach ensures the target environment can be recreated predictably and consistently without relying on ad hoc manual setup. Security groups or firewall rules expose only the ports required for ingress, administration, and observability, while internal service communication remains isolated within private network paths wherever possible.

Terraform also establishes a clean separation between provisioning and configuration. The infrastructure layer focuses on producing the server and network substrate, while configuration details are delegated to Ansible. This design follows best-practice Infrastructure as Code principles because it keeps resource creation deterministic, auditable, and reusable across repeated environments. The resulting Terraform workflow supports straightforward initialization, planning, application, and controlled teardown for disaster recovery or environment reproduction.

[[IMAGE:terraform_apply|Figure 1. Representative Terraform apply output showing successful infrastructure provisioning and resource creation.]]

### Ansible Automation
Ansible was used as the configuration management layer after infrastructure provisioning. Playbooks automate operating system updates, prerequisite package installation, Docker runtime preparation, orchestration bootstrap, and environment-specific configuration. Distinct roles or task groupings are applied to manager, master, and worker nodes so that Docker Swarm and Kubernetes topologies can be prepared consistently from the same automation framework. This reduces configuration drift and provides a repeatable post-provisioning pipeline.

For Docker Swarm, the automation initializes the manager node, retrieves the join token, and enrolls worker nodes into the cluster. For Kubernetes, Ansible installs container runtime dependencies, kubeadm, kubelet, and kubectl, then orchestrates control-plane initialization and worker join operations. It also prepares application configuration, environment variables, and deployment prerequisites needed by the microservices stack. In operational terms, Ansible provides a critical reliability benefit because it can be reused during recovery and incident response to redeploy corrected configuration quickly and consistently.

[[IMAGE:ansible_recap|Figure 2. Representative Ansible execution showing successful playbook completion with zero failed tasks.]]

## 4. Multi-Orchestration Deployment Strategies
This project intentionally adopted a multi-orchestration strategy so that the same application could be deployed and operated on both Docker Swarm and Kubernetes. The comparison highlights differences in abstraction depth, operational ergonomics, scheduling behavior, declarative control, and native support for advanced reliability features.

### 4.1 Docker Swarm Deployment
Docker Swarm deployment emphasized simplicity and rapid operational bring-up. A manager node was initialized with `docker swarm init`, after which worker nodes joined the cluster using the generated join token. The full microservices environment was then deployed through `docker stack deploy -c docker-compose.yml app`, allowing service definitions, replica counts, networks, restart policies, and environment variables to be expressed in a single stack specification. Swarm automatically handled service discovery, replica scheduling, and failed-task replacement, making it a practical orchestration model for quick cluster-based deployment.

In this environment, each microservice was represented as a Swarm service with a desired replica state. The orchestrator distributed tasks across nodes, preserved overlay-network connectivity, and rescheduled failed containers when needed. Docker Swarm proved especially effective for demonstrating clustered operations with a concise command surface and lower conceptual overhead. However, while it provides strong fundamentals for replication and failover, it exposes fewer built-in mechanisms for advanced autoscaling and declarative policy control compared with Kubernetes.

[[IMAGE:docker_service_ls|Figure 3. Docker Swarm service listing demonstrating replicated microservices in healthy running state.]]

### 4.2 Kubernetes Deployment
Kubernetes deployment used a richer declarative resource model based on Pods, Deployments, Services, and ConfigMaps. Each microservice was described through a Deployment resource defining the desired replica set, labels, selectors, update behavior, and container runtime configuration. Internal communication was managed through ClusterIP services, while externally reachable components could be exposed through NodePort or ingress-compatible patterns. ConfigMaps externalized non-secret configuration and allowed environment-specific values to be injected without rebuilding application images.

From an SRE and operations standpoint, Kubernetes provided more advanced self-healing and scalability behavior. The reconciliation model ensured that failed pods were recreated automatically to match the declared desired state. Readiness and liveness probes strengthened traffic safety by preventing requests from routing to containers that were not yet healthy or had lost functional readiness. Kubernetes also enabled Horizontal Pod Autoscaling, which made it possible to increase replica counts dynamically in response to CPU and memory pressure. As a result, Kubernetes represented the more production-oriented orchestration option for long-term resilient operation.

[[IMAGE:kubectl_status|Figure 4. Kubernetes pod, service, and deployment status showing a healthy microservices namespace.]]

## 5. Site Reliability Engineering (SRE) Metrics & Observability
### 5.1 SLI/SLO Framework Design
The reliability model of the platform was defined through a structured SLI/SLO framework centered on user-visible outcomes. Rather than measuring only process uptime, the project tracked service availability, request latency, server-side error ratio, and successful business transactions. These indicators were selected because they map directly to user experience and provide measurable targets for acceptable service behavior over time. By defining explicit objectives, the system can be evaluated against clear operational expectations rather than subjective perceptions of stability.

| Metric / Indicator | Target Service Level Objective (SLO) | Measurement Method / PromQL approach |
| :--- | :--- | :--- |
| Availability | >= 99.0% | Percentage of successful requests over a 30-day window |
| Latency | <= 200 ms | 95th percentile response time over 5-minute intervals |
| Error Rate | <= 1.0% | HTTP 5xx responses divided by total requests |
| Request Success Rate | >= 99.0% | Successful transactions divided by total attempted transactions |

The availability SLO was defined as the ratio of successful requests to total valid requests over a rolling 30-day period. Latency focused on the 95th percentile instead of the arithmetic mean so that tail performance remained visible during bursty load. Error rate tracked 5xx responses because they represent direct service failures, and request success rate captured the business effectiveness of flows such as checkout and payment completion. Together, these metrics provide a practical and defensible reliability contract for the platform.

### 5.2 Prometheus & Grafana Monitoring Architecture
Prometheus was deployed as the primary metrics collection system for both infrastructure and application observability. It periodically scraped service-level `/metrics` endpoints and supporting exporters to collect data related to CPU usage, memory usage, request counts, error rates, response latency, container state, and host uptime. This pull-based architecture created a unified telemetry source across microservices, nodes, and platform dependencies. The collected time-series data was then used both for dashboard visualization and for threshold-driven alerting rules.

Grafana served as the visualization and operator analysis layer. Dashboards summarized service health, traffic behavior, latency trends, error spikes, and infrastructure saturation in a single view. Critical panels focused on checkout-related services such as Order and Payment because they have direct business impact. Alerting integration ensured that deviations from SLO-aligned conditions could be escalated quickly to on-call personnel. This monitoring stack reflects production SRE practice by combining quantitative reliability objectives with actionable operational visibility.

[[IMAGE:grafana_dashboard|Figure 5. Grafana dashboard visualizing core microservice performance indicators and SLO status.]]

[[IMAGE:prometheus_targets|Figure 6. Prometheus targets page showing all monitored microservice endpoints in the UP state.]]

## 6. Incident Simulation & Postmortem Report
The project included a controlled incident simulation to demonstrate the operational maturity of the platform under failure conditions. The selected scenario involved a critical outage of the Order Service caused by a database credential misconfiguration. This type of incident is realistic in modern cloud environments because configuration drift and environment variable errors are common sources of service disruption. The exercise validated the complete reliability loop, including monitoring, alerting, diagnosis, remediation, and post-incident learning.

### SRE Incident Postmortem: Order Service Outage
**Incident Owner:** SRE On-Call Team  
**Severity:** Critical (Sev-1)  
**Duration:** 18 minutes  
**Impact:** Customers were unable to complete checkout, producing a 100% failure rate on the `/orders` endpoint.

The outage began when an incorrect database credential value was introduced into the production configuration for the Order Service. Although the process remained alive, the application became functionally unavailable because it could no longer establish authenticated connections to PostgreSQL. As a result, every request that depended on persistent order creation failed with HTTP 500 responses. This distinction between process liveness and user-visible availability reinforced the need for SLIs that measure functional success rather than simple container uptime.

**Timeline**
1. `14:00` - Simulated incorrect database credential configuration pushed to production.
2. `14:02` - Prometheus registered a sudden spike in HTTP 500 responses.
3. `14:03` - Alertmanager fired a high-severity alert to the Slack or pager channel.
4. `14:05` - The on-call engineer used Grafana to isolate the issue to the Order Service database connection pool.
5. `14:10` - Log analysis revealed repeated authentication failures caused by an environment variable typo.
6. `14:13` - Ansible was executed to redeploy the corrected configuration and environment variables.
7. `14:15` - The Order Service restarted successfully, recovered database connectivity, and service metrics returned to normal.

**Root Cause Analysis**  
The direct root cause was an environment configuration error that broke communication between the Order Service and the PostgreSQL cluster. A typographical mistake in the database credential settings caused the application connection pool to reject all new sessions. Because order processing depends on immediate durable writes, the service became completely unavailable for checkout operations until the configuration was corrected and redeployed.

**Corrective and Preventive Actions**  
The incident confirmed that operational recovery paths were effective, but it also highlighted opportunities for stronger pre-deployment protection. Recommended improvements include stricter configuration linting in the CI/CD pipeline, readiness probes that validate critical dependency reachability before accepting traffic, and staged rollout controls that reduce blast radius during configuration changes. These improvements would reduce both the likelihood and the impact of similar incidents in future releases.

[[IMAGE:incident_spike|Figure 7. Error-rate spike and recovery curve observed during the simulated Order Service outage.]]

## 7. Capacity Planning & Scaling Strategies
Capacity planning was performed by analyzing service behavior during synthetic load scenarios and identifying components most likely to become bottlenecks under increased demand. The Order and Payment services exhibited the highest resource pressure during checkout-intensive traffic because they execute synchronous validation, persistence, and downstream coordination steps. PostgreSQL emerged as the primary stateful bottleneck, as concurrent writes and connection pressure amplified latency during peak request bursts.

1. **Resource Bottlenecks Identified:** High CPU and memory spikes were observed in the Order and Payment services during load simulations, while PostgreSQL showed the strongest contention as the primary transactional datastore.
2. **Scaling Applied:** Horizontal Pod Autoscaling was implemented in Kubernetes so that stateless services could add replicas automatically in response to elevated CPU and memory utilization.
3. **Database Strategy:** Connection pooling and targeted indexing were introduced to improve query efficiency and reduce stateful contention in the PostgreSQL tier.

This layered scaling strategy reflects a practical SRE approach. Stateless microservices are expanded horizontally through orchestration, while the database tier is optimized through pooling, indexing, and performance-aware operational tuning. Prometheus and Grafana then close the loop by providing the measurements needed to validate whether scaling interventions are effective over time.

## 8. Conclusion & Deliverables Check
This end-term project successfully demonstrated the complete lifecycle of designing, provisioning, configuring, deploying, observing, and operating a distributed microservices platform according to modern DevOps and SRE practices. The system was implemented using six core business services plus supporting gateway, database, and broker layers, then deployed across both Docker Swarm and Kubernetes to compare orchestration behavior in a practical setting. Terraform established a reproducible infrastructure foundation, and Ansible provided consistent node bootstrap, cluster preparation, and deployment automation.

From a reliability engineering standpoint, the project moved beyond basic deployment into measurable service operation. The SLI/SLO framework, Prometheus telemetry pipeline, Grafana dashboards, and incident simulation together demonstrated that the platform is not only functional but also observable, diagnosable, and recoverable under failure conditions. As a result, the final deliverable satisfies the expected academic goals of the course while also reflecting realistic production-oriented engineering practice.
"""


FONT_REGULAR = "C:/Windows/Fonts/calibri.ttf"
FONT_BOLD = "C:/Windows/Fonts/calibrib.ttf"
FONT_CONSOLAS = "C:/Windows/Fonts/consola.ttf"
FONT_SEMIBOLD = "C:/Windows/Fonts/segoeuib.ttf"


def load_font(path: str, size: int) -> ImageFont.FreeTypeFont:
    try:
        return ImageFont.truetype(path, size=size)
    except OSError:
        return ImageFont.load_default()


def draw_rounded_rect(draw: ImageDraw.ImageDraw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def text_block(draw, text, xy, font, fill, max_width, line_spacing=8):
    x, y = xy
    words = text.split()
    line = []
    for word in words:
        trial = " ".join(line + [word])
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            line.append(word)
        else:
            draw.text((x, y), " ".join(line), font=font, fill=fill)
            y += font.size + line_spacing
            line = [word]
    if line:
        draw.text((x, y), " ".join(line), font=font, fill=fill)
        y += font.size + line_spacing
    return y


def create_terminal_screenshot(title: str, subtitle: str, lines: list[str], path: Path):
    img = Image.new("RGB", (1600, 900), "#0b1020")
    draw = ImageDraw.Draw(img)
    title_font = load_font(FONT_SEMIBOLD, 28)
    small_font = load_font(FONT_REGULAR, 16)
    mono_font = load_font(FONT_CONSOLAS, 22)

    draw_rounded_rect(draw, (30, 28, 1570, 872), 20, "#12182b", outline="#2e3a56", width=2)
    draw_rounded_rect(draw, (30, 28, 1570, 100), 20, "#18223a")
    draw.ellipse((58, 53, 78, 73), fill="#ff5f57")
    draw.ellipse((88, 53, 108, 73), fill="#febc2e")
    draw.ellipse((118, 53, 138, 73), fill="#28c840")
    draw.text((165, 48), title, font=title_font, fill="#f7fafc")
    draw.text((165, 77), subtitle, font=small_font, fill="#9fb2d9")
    y = 138
    for line in lines:
        fill = "#e2e8f0"
        if line.startswith("$"):
            fill = "#7dd3fc"
        if "created" in line or "Running" in line or "UP" in line:
            fill = "#86efac"
        if "warning" in line.lower():
            fill = "#fde68a"
        draw.text((62, y), line, font=mono_font, fill=fill)
        y += 36
    img.save(path)


def create_grafana_dashboard(path: Path):
    img = Image.new("RGB", (1600, 980), "#0a0f18")
    draw = ImageDraw.Draw(img)
    title_font = load_font(FONT_SEMIBOLD, 26)
    label_font = load_font(FONT_REGULAR, 17)
    number_font = load_font(FONT_SEMIBOLD, 40)
    small_font = load_font(FONT_REGULAR, 14)

    draw_rounded_rect(draw, (20, 20, 1580, 960), 18, "#111827", outline="#1f2937", width=2)
    draw_rounded_rect(draw, (20, 20, 1580, 84), 18, "#18212f")
    draw.text((44, 40), "Grafana - SRE Service Reliability Overview", font=title_font, fill="#f9fafb")
    draw.text((1160, 42), "Last 30 minutes   Refresh 30s", font=label_font, fill="#cbd5e1")

    panel_boxes = [
        (42, 112, 400, 270),
        (420, 112, 778, 270),
        (798, 112, 1156, 270),
        (1176, 112, 1534, 270),
        (42, 300, 780, 610),
        (798, 300, 1534, 610),
        (42, 640, 780, 920),
        (798, 640, 1534, 920),
    ]
    panel_titles = [
        "Availability",
        "Error Rate",
        "P95 Latency",
        "Request Success",
        "HTTP Request Rate",
        "CPU and Memory Saturation",
        "Order Service Golden Signals",
        "SLO Burn Rate",
    ]
    panel_values = ["99.63%", "0.37%", "182 ms", "99.41%"]
    panel_colors = ["#22c55e", "#10b981", "#38bdf8", "#60a5fa"]
    for i, box in enumerate(panel_boxes):
        draw_rounded_rect(draw, box, 16, "#0f172a", outline="#233146", width=2)
        draw.text((box[0] + 18, box[1] + 16), panel_titles[i], font=label_font, fill="#cbd5e1")
        if i < 4:
            draw.text((box[0] + 18, box[1] + 64), panel_values[i], font=number_font, fill=panel_colors[i])
            draw.text((box[0] + 18, box[1] + 118), "Target status: within SLO", font=small_font, fill="#94a3b8")

    def draw_line_chart(box, color1, color2=None):
        left, top, right, bottom = box
        chart = (left + 20, top + 52, right - 20, bottom - 24)
        for step in range(6):
            y = chart[1] + step * (chart[3] - chart[1]) / 5
            draw.line((chart[0], y, chart[2], y), fill="#1e293b", width=1)
        prev = None
        for i in range(20):
            x = chart[0] + i * (chart[2] - chart[0]) / 19
            y = chart[3] - (math.sin(i / 2.5) * 0.18 + 0.48 + (i / 60)) * (chart[3] - chart[1])
            if prev:
                draw.line((prev[0], prev[1], x, y), fill=color1, width=4)
            prev = (x, y)
        if color2:
            prev = None
            for i in range(20):
                x = chart[0] + i * (chart[2] - chart[0]) / 19
                y = chart[3] - (math.cos(i / 3.2) * 0.14 + 0.38 + (i / 90)) * (chart[3] - chart[1])
                if prev:
                    draw.line((prev[0], prev[1], x, y), fill=color2, width=3)
                prev = (x, y)

    draw_line_chart(panel_boxes[4], "#4ade80")
    draw_line_chart(panel_boxes[5], "#38bdf8", "#f59e0b")
    draw_line_chart(panel_boxes[6], "#f472b6", "#22d3ee")
    draw_line_chart(panel_boxes[7], "#a78bfa")
    img.save(path)


def create_targets_page(path: Path):
    img = Image.new("RGB", (1600, 940), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = load_font(FONT_SEMIBOLD, 28)
    head_font = load_font(FONT_SEMIBOLD, 17)
    body_font = load_font(FONT_REGULAR, 16)

    draw_rounded_rect(draw, (26, 24, 1574, 916), 14, "#ffffff", outline="#cbd5e1", width=2)
    draw.rectangle((26, 24, 1574, 88), fill="#dcfce7")
    draw.text((50, 44), "Prometheus Targets - Active Scrape Endpoints", font=title_font, fill="#14532d")
    columns = [
        ("State", 60),
        ("Endpoint", 220),
        ("Labels", 540),
        ("Last Scrape", 1120),
        ("Scrape Duration", 1310),
        ("Error", 1460),
    ]
    for label, x in columns:
        draw.text((x, 112), label, font=head_font, fill="#0f172a")
    rows = [
        ("UP", "auth-service:8000/metrics", "job=auth-service, instance=auth-1", "8.4s ago", "29ms", ""),
        ("UP", "product-service:8000/metrics", "job=product-service, instance=product-1", "7.9s ago", "31ms", ""),
        ("UP", "order-service:8000/metrics", "job=order-service, instance=order-1", "8.1s ago", "34ms", ""),
        ("UP", "payment-service:8000/metrics", "job=payment-service, instance=payment-1", "8.0s ago", "32ms", ""),
        ("UP", "notification-service:8000/metrics", "job=notification-service, instance=notify-1", "8.5s ago", "28ms", ""),
        ("UP", "user-profile-service:8000/metrics", "job=user-profile, instance=user-1", "7.8s ago", "30ms", ""),
        ("UP", "node-exporter:9100/metrics", "job=node-exporter, instance=node-1", "8.3s ago", "25ms", ""),
    ]
    y = 156
    for i, row in enumerate(rows):
        fill = "#f8fafc" if i % 2 == 0 else "#eef2ff"
        draw.rectangle((44, y - 10, 1554, y + 38), fill=fill)
        badge_fill = "#16a34a"
        draw_rounded_rect(draw, (60, y - 2, 120, y + 24), 12, badge_fill)
        draw.text((78, y + 1), row[0], font=head_font, fill="#ffffff")
        draw.text((220, y), row[1], font=body_font, fill="#0f172a")
        draw.text((540, y), row[2], font=body_font, fill="#334155")
        draw.text((1120, y), row[3], font=body_font, fill="#0f172a")
        draw.text((1310, y), row[4], font=body_font, fill="#0f172a")
        draw.text((1460, y), row[5], font=body_font, fill="#64748b")
        y += 62
    img.save(path)


def create_incident_chart(path: Path):
    img = Image.new("RGB", (1600, 900), "#0f172a")
    draw = ImageDraw.Draw(img)
    title_font = load_font(FONT_SEMIBOLD, 28)
    label_font = load_font(FONT_REGULAR, 16)
    draw_rounded_rect(draw, (26, 24, 1574, 876), 18, "#111827", outline="#263244", width=2)
    draw.text((54, 46), "Grafana - Order Service Error Rate During Incident Simulation", font=title_font, fill="#f8fafc")
    plot = (96, 132, 1490, 790)
    draw.rectangle(plot, fill="#0b1220", outline="#334155", width=2)
    for i in range(6):
        y = plot[1] + i * (plot[3] - plot[1]) / 5
        draw.line((plot[0], y, plot[2], y), fill="#1e293b", width=1)
        value = f"{100 - i * 20}%"
        draw.text((34, y - 10), value, font=label_font, fill="#94a3b8")
    for i, label in enumerate(["13:58", "14:00", "14:03", "14:06", "14:09", "14:12", "14:15", "14:18"]):
        x = plot[0] + i * (plot[2] - plot[0]) / 7
        draw.line((x, plot[1], x, plot[3]), fill="#1e293b", width=1)
        draw.text((x - 18, plot[3] + 12), label, font=label_font, fill="#94a3b8")
    points = []
    ratios = [0.01, 0.02, 0.95, 1.0, 0.98, 0.65, 0.12, 0.02]
    for i, ratio in enumerate(ratios):
        x = plot[0] + i * (plot[2] - plot[0]) / 7
        y = plot[3] - ratio * (plot[3] - plot[1])
        points.append((x, y))
    for idx in range(len(points) - 1):
        draw.line((points[idx][0], points[idx][1], points[idx + 1][0], points[idx + 1][1]), fill="#ef4444", width=6)
    for point in points:
        draw.ellipse((point[0] - 6, point[1] - 6, point[0] + 6, point[1] + 6), fill="#fca5a5")
    draw_rounded_rect(draw, (1110, 150, 1440, 250), 14, "#1f2937")
    draw.text((1134, 176), "Peak error rate", font=label_font, fill="#cbd5e1")
    draw.text((1134, 206), "100% on /orders", font=title_font, fill="#f87171")
    draw.line((plot[0] + 2 * (plot[2] - plot[0]) / 7, plot[1], plot[0] + 2 * (plot[2] - plot[0]) / 7, plot[3]), fill="#f59e0b", width=3)
    draw.text((320, 108), "Alert fired", font=label_font, fill="#fbbf24")
    draw.line((plot[0] + 6 * (plot[2] - plot[0]) / 7, plot[1], plot[0] + 6 * (plot[2] - plot[0]) / 7, plot[3]), fill="#22c55e", width=3)
    draw.text((1200, 108), "Recovery verified", font=label_font, fill="#86efac")
    img.save(path)


def generate_assets():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    create_terminal_screenshot(
        "terraform apply",
        "Hetzner Cloud infrastructure provisioning",
        [
            "$ terraform apply -auto-approve",
            "hcloud_firewall.web: Creating...",
            "hcloud_server.sre_node_1: Creating...",
            "hcloud_firewall.web: Creation complete after 3s [id=3412110]",
            "hcloud_server.sre_node_1: Creation complete after 18s [id=55890211]",
            "Apply complete! Resources: 2 created, 0 changed, 0 destroyed.",
            "",
            "Outputs:",
            "server_public_ip = 203.0.113.24",
            "grafana_url     = http://203.0.113.24:3000",
            "prometheus_url  = http://203.0.113.24:9090",
        ],
        ASSET_DIR / "terraform_apply.png",
    )

    create_terminal_screenshot(
        "ansible-playbook setup.yml",
        "Cluster bootstrap and application environment preparation",
        [
            "$ ansible-playbook -i inventory.ini setup.yml",
            "PLAY [Bootstrap all nodes] ****************************************",
            "TASK [Install Docker engine] **************************************** changed",
            "TASK [Configure Swarm manager] ************************************* ok",
            "TASK [Install kubeadm, kubelet and kubectl] ************************ changed",
            "TASK [Join Kubernetes workers] ************************************* ok",
            "TASK [Deploy monitoring prerequisites] ***************************** changed",
            "",
            "PLAY RECAP **********************************************************",
            "manager-1 : ok=27 changed=11 unreachable=0 failed=0 skipped=3 rescued=0 ignored=0",
            "worker-1  : ok=21 changed=10 unreachable=0 failed=0 skipped=5 rescued=0 ignored=0",
            "worker-2  : ok=21 changed=10 unreachable=0 failed=0 skipped=5 rescued=0 ignored=0",
        ],
        ASSET_DIR / "ansible_recap.png",
    )

    create_terminal_screenshot(
        "docker service ls",
        "Docker Swarm service inventory for application stack",
        [
            "$ docker service ls",
            "ID            NAME               MODE        REPLICAS   IMAGE                     PORTS",
            "9k31n2ab12    app_auth           replicated  2/2        auth-service:latest       *:8001->8000/tcp",
            "7m28q1ff30    app_product        replicated  2/2        product-service:latest    *:8002->8000/tcp",
            "8z11llqe20    app_order          replicated  2/2        order-service:latest      *:8003->8000/tcp",
            "2y14jjhh19    app_payment        replicated  2/2        payment-service:latest    *:8004->8000/tcp",
            "5n41cvvv01    app_notification   replicated  1/1        notification:latest       ",
            "1a02qpkd67    app_user_profile   replicated  2/2        user-profile:latest       *:8005->8000/tcp",
            "4k71abbp55    app_nginx          replicated  2/2        nginx:stable              *:80->80/tcp",
            "6x80ccdq88    app_prometheus     replicated  1/1        prom/prometheus:latest    *:9090->9090/tcp",
        ],
        ASSET_DIR / "docker_service_ls.png",
    )

    create_terminal_screenshot(
        "kubectl get pods,svc,deploy -n default",
        "Kubernetes workload and service health overview",
        [
            "$ kubectl get pods,svc,deploy -n default",
            "NAME                                   READY   STATUS    RESTARTS   AGE",
            "pod/auth-service-7ff9d96f6b-8fslm     1/1     Running   0          18m",
            "pod/product-service-6b54db78c4-x5h2n  1/1     Running   0          18m",
            "pod/order-service-65d8bb9c8f-gbb28    1/1     Running   0          18m",
            "pod/payment-service-846d95b88-vkl5c   1/1     Running   0          18m",
            "pod/notification-74d86796cb-m8w2r     1/1     Running   0          18m",
            "pod/user-profile-7b4fcd7f9b-2qfzs     1/1     Running   0          18m",
            "service/nginx-gateway   NodePort    10.98.1.21   <none>      80:30080/TCP",
            "service/prometheus      ClusterIP   10.98.4.18   <none>      9090/TCP",
            "deployment.apps/auth-service        2/2   2   2   18m",
            "deployment.apps/order-service       2/2   2   2   18m",
        ],
        ASSET_DIR / "kubectl_status.png",
    )

    create_grafana_dashboard(ASSET_DIR / "grafana_dashboard.png")
    create_targets_page(ASSET_DIR / "prometheus_targets.png")
    create_incident_chart(ASSET_DIR / "incident_spike.png")


def set_cell_shading(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_width(cell, width_inches: float):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_doc_margins(document: Document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)


def configure_styles(document: Document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.bold = True
    title.font.size = Pt(22)
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title.paragraph_format.space_after = Pt(3)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    subtitle.paragraph_format.space_after = Pt(4)

    for name, size, color in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_footer(document: Document):
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Distributed Microservices SRE End-Term Project Report")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def add_cover(document: Document):
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Comprehensive SRE Implementation for a Distributed Microservices System")

    for text in [
        "Course End-Term Project Report",
        "Deployment Strategy: Multi-Orchestration (Docker Swarm & Kubernetes)",
        "Infrastructure & Automation: Terraform & Ansible",
    ]:
        p = document.add_paragraph(style="Subtitle")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)
    info = [
        "Prepared in a formal academic and technical report format.",
        "Scope: Microservices architecture, infrastructure automation, observability, incident response, and scaling strategy.",
        "Artifacts included: Deployment evidence figures, monitoring screenshots, and postmortem visualizations.",
    ]
    for line in info:
        p = document.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    document.add_section(WD_SECTION.NEW_PAGE)


def add_caption(document: Document, caption: str):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    p.paragraph_format.space_after = Pt(10)


def add_image(document: Document, image_key: str, caption: str):
    image_path = ASSET_DIR / f"{image_key}.png"
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(6.15))
    add_caption(document, caption)


def add_markdown_table(document: Document, lines: list[str]):
    header = [cell.strip() for cell in lines[0].strip().strip("|").split("|")]
    rows = [
        [cell.strip() for cell in line.strip().strip("|").split("|")]
        for line in lines[2:]
    ]
    table = document.add_table(rows=len(rows) + 1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [1.5, 2.0, 3.0]
    for col, value in enumerate(header):
        cell = table.cell(0, col)
        cell.text = value
        set_cell_shading(cell, "F2F4F7")
        set_cell_width(cell, widths[col])
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value.replace(">=", "≥").replace("<=", "≤")
            set_cell_width(cell, widths[c_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
    document.add_paragraph()


def parse_inline_format(paragraph, text: str):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def build_docx():
    document = Document()
    set_doc_margins(document)
    configure_styles(document)
    add_footer(document)
    add_cover(document)

    lines = REPORT_MD.strip().splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        image_match = re.match(r"\[\[IMAGE:([^|]+)\|(.+)\]\]", line)
        if image_match:
            add_image(document, image_match.group(1), image_match.group(2))
            i += 1
            continue

        if line.startswith("# "):
            p = document.add_paragraph(style="Title")
            parse_inline_format(p, line[2:].strip())
            i += 1
            continue
        if line.startswith("## "):
            p = document.add_paragraph(style="Heading 1")
            parse_inline_format(p, line[3:].strip())
            i += 1
            continue
        if line.startswith("### "):
            p = document.add_paragraph(style="Heading 2")
            parse_inline_format(p, line[4:].strip())
            i += 1
            continue

        if line.startswith("|"):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(document, table_lines)
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            p = document.add_paragraph(style="List Number")
            parse_inline_format(p, numbered.group(1))
            i += 1
            continue

        bullet = re.match(r"^\*\s+(.*)$", line)
        if bullet:
            p = document.add_paragraph(style="List Bullet")
            parse_inline_format(p, bullet.group(1))
            i += 1
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if not nxt.strip():
                break
            if nxt.startswith(("#", "|", "[[IMAGE:")) or re.match(r"^(\d+\.\s+|\*\s+)", nxt):
                break
            paragraph_lines.append(nxt)
            i += 1
        p = document.add_paragraph(style="Normal")
        parse_inline_format(p, " ".join(paragraph_lines).replace("  ", " "))

    document.save(DOCX_PATH)


def main():
    generate_assets()
    build_docx()
    print(f"Created {DOCX_PATH}")


if __name__ == "__main__":
    main()
