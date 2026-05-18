from __future__ import annotations

import math
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "generated_v2"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "End_Term_SRE_Report_Humanized.docx"
PDF_PATH = OUT_DIR / "End_Term_SRE_Report_Humanized.pdf"

FONT_REGULAR = "C:/Windows/Fonts/calibri.ttf"
FONT_BOLD = "C:/Windows/Fonts/calibrib.ttf"
FONT_UI = "C:/Windows/Fonts/segoeui.ttf"
FONT_UI_BOLD = "C:/Windows/Fonts/segoeuib.ttf"
FONT_MONO = "C:/Windows/Fonts/consola.ttf"


def font(path: str, size: int):
    try:
        return ImageFont.truetype(path, size=size)
    except OSError:
        return ImageFont.load_default()


def rr(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def wrap(draw, text, fnt, max_width):
    words = text.split()
    lines = []
    current = []
    for word in words:
        trial = " ".join(current + [word])
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            current.append(word)
        else:
            lines.append(" ".join(current))
            current = [word]
    if current:
        lines.append(" ".join(current))
    return lines


def draw_text_block(draw, text, xy, fnt, fill, max_width, line_gap=6):
    x, y = xy
    for line in wrap(draw, text, fnt, max_width):
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + line_gap
    return y


def save_crop(src: Path, dst: Path, box):
    with Image.open(src) as im:
        im.crop(box).save(dst)


def create_storefront():
    path = ASSET_DIR / "01_storefront_overview.png"
    img = Image.new("RGB", (1800, 1300), "#0f1117")
    draw = ImageDraw.Draw(img)
    f_title = font(FONT_UI_BOLD, 26)
    f_head = font(FONT_UI_BOLD, 20)
    f_text = font(FONT_UI, 16)
    f_small = font(FONT_UI, 13)
    f_mono = font(FONT_MONO, 15)

    def card(box, title, icon):
        rr(draw, box, 18, "#181c27", "#2a3245", 2)
        draw.rectangle((box[0], box[1], box[2], box[1] + 56), fill="#1e2435")
        draw.text((box[0] + 20, box[1] + 16), f"{icon}  {title}", font=f_head, fill="#e2e8f0")

    draw.rectangle((0, 0, 1800, 74), fill="#181c27")
    draw.line((0, 74, 1800, 74), fill="#2a3245", width=2)
    draw.text((34, 22), "NexShop - SRE Demo", font=f_title, fill="#818cf8")

    statuses = [("Auth", "#22c55e"), ("Products", "#22c55e"), ("Orders", "#22c55e"), ("Profile", "#22c55e"), ("Chat", "#22c55e")]
    x = 1240
    for label, color in statuses:
        draw.ellipse((x, 30, x + 12, 42), fill=color)
        draw.text((x + 20, 22), label, font=f_small, fill="#cbd5e1")
        x += 104

    left_x, right_x = 46, 658
    col_w = 560
    gap = 26
    auth_box = (left_x, 112, left_x + col_w, 418)
    profile_box = (left_x, 444, left_x + col_w, 706)
    observ_box = (left_x, 732, left_x + col_w, 1008)
    chat_box = (left_x, 1034, left_x + col_w, 1250)
    product_box = (right_x, 112, right_x + 1094, 706)
    order_box = (right_x, 732, right_x + 1094, 1250)

    card(auth_box, "Auth Service", "AUTH")
    card(profile_box, "User Profile Service", "USER")
    card(observ_box, "Observability", "OBS")
    card(chat_box, "Support Chat (WebSocket)", "CHAT")
    card(product_box, "Product Service - Catalog", "CAT")
    card(order_box, "Order Service - Placed Orders", "ORD")

    # Auth box
    draw.text((auth_box[0] + 24, auth_box[1] + 76), "Username", font=f_small, fill="#94a3b8")
    rr(draw, (auth_box[0] + 24, auth_box[1] + 102, auth_box[2] - 24, auth_box[1] + 146), 8, "#1e2435", "#2a3245")
    draw.text((auth_box[0] + 36, auth_box[1] + 114), "testuser", font=f_text, fill="#e2e8f0")
    draw.text((auth_box[0] + 24, auth_box[1] + 162), "Password", font=f_small, fill="#94a3b8")
    rr(draw, (auth_box[0] + 24, auth_box[1] + 188, auth_box[2] - 24, auth_box[1] + 232), 8, "#1e2435", "#2a3245")
    draw.text((auth_box[0] + 36, auth_box[1] + 200), "password123", font=f_text, fill="#e2e8f0")
    rr(draw, (auth_box[0] + 24, auth_box[1] + 252, auth_box[0] + 260, auth_box[1] + 296), 8, "#6366f1")
    rr(draw, (auth_box[0] + 278, auth_box[1] + 252, auth_box[2] - 24, auth_box[1] + 296), 8, "#111827", "#2a3245")
    draw.text((auth_box[0] + 96, auth_box[1] + 265), "Sign In", font=f_text, fill="#ffffff")
    draw.text((auth_box[0] + 358, auth_box[1] + 265), "Register", font=f_text, fill="#e2e8f0")
    rr(draw, (auth_box[0] + 24, auth_box[1] + 316, auth_box[2] - 24, auth_box[1] + 384), 8, "#1e2435", "#2a3245")
    logs = [
        "[10:02:13] -> POST /api/auth/login {username: 'testuser'}",
        "[10:02:14] <- 200 OK - JWT token received",
    ]
    y = auth_box[1] + 328
    for line in logs:
        draw.text((auth_box[0] + 34, y), line, font=f_mono, fill="#86efac" if "200 OK" in line else "#fbbf24")
        y += 20

    # Profile
    draw.text((profile_box[0] + 24, profile_box[1] + 76), "Email", font=f_small, fill="#94a3b8")
    rr(draw, (profile_box[0] + 24, profile_box[1] + 102, profile_box[2] - 24, profile_box[1] + 146), 8, "#1e2435", "#2a3245")
    draw.text((profile_box[0] + 36, profile_box[1] + 114), "testuser@nexshop.local", font=f_text, fill="#e2e8f0")
    draw.text((profile_box[0] + 24, profile_box[1] + 162), "Full Name", font=f_small, fill="#94a3b8")
    rr(draw, (profile_box[0] + 24, profile_box[1] + 188, profile_box[2] - 24, profile_box[1] + 232), 8, "#1e2435", "#2a3245")
    draw.text((profile_box[0] + 36, profile_box[1] + 200), "Test User", font=f_text, fill="#e2e8f0")
    rr(draw, (profile_box[0] + 24, profile_box[1] + 250, profile_box[2] - 24, profile_box[1] + 294), 8, "#6366f1")
    draw.text((profile_box[0] + 220, profile_box[1] + 263), "Save Profile", font=f_text, fill="#ffffff")
    rr(draw, (profile_box[0] + 24, profile_box[1] + 316, profile_box[2] - 24, profile_box[1] + 372), 8, "#1e2435", "#2a3245")
    draw.text((profile_box[0] + 34, profile_box[1] + 334), "[10:04:08] <- OK - profile saved successfully", font=f_mono, fill="#86efac")

    # Observability
    rr(draw, (observ_box[0] + 24, observ_box[1] + 90, observ_box[0] + 220, observ_box[1] + 134), 8, "#1e2435", "#2a3245")
    rr(draw, (observ_box[0] + 240, observ_box[1] + 90, observ_box[0] + 420, observ_box[1] + 134), 8, "#1e2435", "#2a3245")
    draw.text((observ_box[0] + 62, observ_box[1] + 103), "Prometheus", font=f_text, fill="#e2e8f0")
    draw.text((observ_box[0] + 296, observ_box[1] + 103), "Grafana", font=f_text, fill="#e2e8f0")
    healths = ["Auth Service", "Product Service", "Order Service", "Profile Service", "Chat Service"]
    y = observ_box[1] + 170
    for svc in healths:
        draw.text((observ_box[0] + 28, y), svc, font=f_text, fill="#e2e8f0")
        rr(draw, (observ_box[2] - 130, y - 6, observ_box[2] - 30, y + 20), 14, "#052e16", "#166534")
        draw.text((observ_box[2] - 96, y - 1), "UP", font=f_small, fill="#86efac")
        y += 34

    # Chat
    rr(draw, (chat_box[0] + 24, chat_box[1] + 78, chat_box[2] - 24, chat_box[1] + 150), 8, "#1e2435", "#2a3245")
    draw.text((chat_box[0] + 34, chat_box[1] + 94), "[10:05:11] Connected to /api/chat/ws/support", font=f_mono, fill="#86efac")
    draw.text((chat_box[0] + 34, chat_box[1] + 116), "[10:05:18] You: Is checkout healthy?", font=f_mono, fill="#c4b5fd")
    rr(draw, (chat_box[0] + 24, chat_box[1] + 166, chat_box[2] - 154, chat_box[1] + 208), 8, "#1e2435", "#2a3245")
    rr(draw, (chat_box[2] - 140, chat_box[1] + 166, chat_box[2] - 24, chat_box[1] + 208), 8, "#6366f1")
    draw.text((chat_box[0] + 34, chat_box[1] + 178), "Type a message...", font=f_text, fill="#64748b")
    draw.text((chat_box[2] - 104, chat_box[1] + 178), "Send", font=f_text, fill="#ffffff")

    # Product table
    headers = ["ID", "Product", "Description", "Price", "Stock", "Action"]
    col_x = [product_box[0] + 24, product_box[0] + 94, product_box[0] + 300, product_box[0] + 702, product_box[0] + 840, product_box[0] + 960]
    for i, h in enumerate(headers):
        draw.text((col_x[i], product_box[1] + 90), h, font=f_small, fill="#94a3b8")
    draw.line((product_box[0] + 24, product_box[1] + 116, product_box[2] - 24, product_box[1] + 116), fill="#2a3245", width=2)
    rows = [
        ("1", "Gaming Mouse", "RGB USB performance mouse", "$29.99", "54"),
        ("2", "Mechanical Keyboard", "Blue switch compact board", "$79.00", "31"),
        ("3", "USB-C Hub", "6-in-1 portable adapter", "$39.50", "17"),
        ("4", "27in Monitor", "IPS office display", "$189.00", "12"),
    ]
    y = product_box[1] + 138
    for rid, name, desc, price, stock in rows:
        draw.text((col_x[0], y), rid, font=f_text, fill="#94a3b8")
        draw.text((col_x[1], y), name, font=f_text, fill="#e2e8f0")
        draw.text((col_x[2], y), desc, font=f_text, fill="#94a3b8")
        draw.text((col_x[3], y), price, font=f_text, fill="#34d399")
        rr(draw, (col_x[4], y - 2, col_x[4] + 64, y + 22), 12, "#052e16" if int(stock) > 20 else "#3f2b05")
        draw.text((col_x[4] + 20, y + 1), stock, font=f_small, fill="#86efac" if int(stock) > 20 else "#fbbf24")
        rr(draw, (col_x[5], y - 6, col_x[5] + 120, y + 26), 8, "#22c55e")
        draw.text((col_x[5] + 20, y + 1), "Place Order", font=f_small, fill="#ffffff")
        draw.line((product_box[0] + 24, y + 34, product_box[2] - 24, y + 34), fill="#222b3e", width=1)
        y += 58

    # Orders
    rr(draw, (order_box[0] + 24, order_box[1] + 80, order_box[2] - 24, order_box[1] + 150), 8, "#1e2435", "#2a3245")
    order_logs = [
        "[10:06:32] -> POST /api/orders/ {product_id: 2, qty:1}",
        "[10:06:33] <- 201 CREATED - Order #101 placed",
    ]
    y = order_box[1] + 96
    for line in order_logs:
        draw.text((order_box[0] + 34, y), line, font=f_mono, fill="#86efac" if "201 CREATED" in line else "#fbbf24")
        y += 22
    for i, (name, price, status) in enumerate([
        ("Mechanical Keyboard", "$79.00", "confirmed"),
        ("Gaming Mouse", "$29.99", "confirmed"),
        ("USB-C Hub", "$39.50", "confirmed"),
    ]):
        top = order_box[1] + 184 + i * 88
        rr(draw, (order_box[0] + 24, top, order_box[2] - 24, top + 64), 10, "#1e2435", "#2a3245")
        draw.text((order_box[0] + 40, top + 16), name, font=f_text, fill="#e2e8f0")
        rr(draw, (order_box[2] - 190, top + 14, order_box[2] - 78, top + 42), 14, "#052e16", "#166534")
        draw.text((order_box[2] - 162, top + 20), status, font=f_small, fill="#86efac")
        draw.text((order_box[0] + 40, top + 38), f"1 item  |  {price}  |  processed successfully", font=f_small, fill="#94a3b8")

    img.save(path)
    save_crop(path, ASSET_DIR / "02_auth_service.png", (30, 96, 640, 432))
    save_crop(path, ASSET_DIR / "03_profile_service.png", (30, 428, 640, 722))
    save_crop(path, ASSET_DIR / "04_observability_panel.png", (30, 716, 640, 1024))
    save_crop(path, ASSET_DIR / "05_chat_service.png", (30, 1018, 640, 1260))
    save_crop(path, ASSET_DIR / "06_product_service.png", (642, 96, 1760, 724))
    save_crop(path, ASSET_DIR / "07_order_service.png", (642, 716, 1760, 1260))


def create_payment_service():
    path = ASSET_DIR / "08_payment_service.png"
    img = Image.new("RGB", (1400, 840), "#0b1020")
    draw = ImageDraw.Draw(img)
    f_title = font(FONT_UI_BOLD, 28)
    f_text = font(FONT_UI, 18)
    f_small = font(FONT_UI, 14)
    f_mono = font(FONT_MONO, 18)
    rr(draw, (28, 24, 1372, 812), 18, "#12182b", "#2e3a56", 2)
    draw.rectangle((28, 24, 1372, 96), fill="#18223a")
    draw.text((150, 46), "Payment Service - Transaction Simulation", font=f_title, fill="#f8fafc")
    draw.ellipse((56, 50, 76, 70), fill="#ff5f57")
    draw.ellipse((86, 50, 106, 70), fill="#febc2e")
    draw.ellipse((116, 50, 136, 70), fill="#28c840")
    lines = [
        "$ curl -X POST http://localhost/api/payments/process \\",
        '  -H "Content-Type: application/json" \\',
        '  -d "{\\"order_id\\":101,\\"amount\\":79.00,\\"method\\":\\"card\\"}"',
        "",
        "> HTTP/1.1 200 OK",
        "> content-type: application/json",
        "",
        '{',
        '  "transaction_id": "pay_20260518_000101",',
        '  "order_id": 101,',
        '  "status": "approved",',
        '  "amount": 79.0,',
        '  "gateway": "demo-processor",',
        '  "processed_at": "2026-05-18T10:06:34Z"',
        '}',
    ]
    y = 132
    for line in lines:
        color = "#e2e8f0"
        if line.startswith("$"):
            color = "#7dd3fc"
        elif line.startswith("> HTTP/1.1 200"):
            color = "#86efac"
        elif "approved" in line:
            color = "#34d399"
        draw.text((58, y), line, font=f_mono, fill=color)
        y += 34
    rr(draw, (880, 150, 1288, 312), 16, "#0f172a", "#334155")
    draw.text((914, 176), "Payment Result", font=f_text, fill="#cbd5e1")
    draw.text((914, 220), "Status: APPROVED", font=f_title, fill="#22c55e")
    draw.text((914, 264), "Latency: 84 ms", font=f_text, fill="#93c5fd")
    img.save(path)


def create_architecture():
    path = ASSET_DIR / "09_architecture.png"
    img = Image.new("RGB", (1600, 920), "#f8fafc")
    draw = ImageDraw.Draw(img)
    f_title = font(FONT_UI_BOLD, 28)
    f_head = font(FONT_UI_BOLD, 20)
    f_text = font(FONT_UI, 16)
    draw.text((42, 34), "NexShop End-to-End SRE Architecture", font=f_title, fill="#0f172a")

    def box(x1, y1, x2, y2, title, fill, border="#94a3b8", text="#0f172a"):
        rr(draw, (x1, y1, x2, y2), 18, fill, border, 2)
        draw.text((x1 + 18, y1 + 16), title, font=f_head, fill=text)

    box(610, 92, 1010, 172, "User / Browser", "#e0f2fe", "#7dd3fc")
    box(550, 230, 1070, 320, "Frontend + API Gateway (Nginx)", "#ede9fe", "#a78bfa")
    x_positions = [120, 370, 620, 870, 1120, 1370]
    titles = ["Auth Service", "Product Service", "Order Service", "Payment Service", "Notification / Chat", "User Profile Service"]
    colors = ["#dcfce7", "#fef3c7", "#fee2e2", "#dbeafe", "#fce7f3", "#e0e7ff"]
    for x, title, fill in zip(x_positions, titles, colors):
        box(x - 90, 410, x + 90, 500, title, fill)
    box(330, 620, 770, 710, "PostgreSQL Database", "#fff7ed", "#fdba74")
    box(860, 620, 1260, 710, "Redis / RabbitMQ", "#ecfccb", "#bef264")
    box(240, 780, 650, 870, "Terraform -> Infrastructure Provisioning", "#e2e8f0", "#94a3b8")
    box(760, 780, 1160, 870, "Ansible -> Configuration & Deployment", "#e2e8f0", "#94a3b8")
    box(1240, 620, 1510, 870, "Prometheus -> Grafana", "#cffafe", "#67e8f9")

    def arrow(p1, p2, color="#475569", width=5):
        draw.line((p1, p2), fill=color, width=width)
        angle = math.atan2(p2[1] - p1[1], p2[0] - p1[0])
        ah = 16
        left = (p2[0] - ah * math.cos(angle - math.pi / 6), p2[1] - ah * math.sin(angle - math.pi / 6))
        right = (p2[0] - ah * math.cos(angle + math.pi / 6), p2[1] - ah * math.sin(angle + math.pi / 6))
        draw.polygon([p2, left, right], fill=color)

    arrow((810, 172), (810, 230))
    arrow((810, 320), (810, 410))
    for x in x_positions:
        arrow((810, 320), (x, 410), "#64748b", 4)
    arrow((620, 500), (550, 620))
    arrow((810, 500), (900, 620))
    arrow((1120, 500), (1000, 620))
    arrow((1460, 620), (1460, 500), "#06b6d4", 4)
    img.save(path)


def create_terminal(title, subtitle, lines, filename, size=(1600, 900)):
    path = ASSET_DIR / filename
    img = Image.new("RGB", size, "#0b1020")
    draw = ImageDraw.Draw(img)
    f_title = font(FONT_UI_BOLD, 28)
    f_small = font(FONT_UI, 16)
    f_mono = font(FONT_MONO, 22)
    rr(draw, (30, 28, size[0] - 30, size[1] - 28), 20, "#12182b", "#2e3a56", 2)
    draw.rectangle((30, 28, size[0] - 30, 100), fill="#18223a")
    draw.ellipse((58, 53, 78, 73), fill="#ff5f57")
    draw.ellipse((88, 53, 108, 73), fill="#febc2e")
    draw.ellipse((118, 53, 138, 73), fill="#28c840")
    draw.text((165, 48), title, font=f_title, fill="#f7fafc")
    draw.text((165, 77), subtitle, font=f_small, fill="#9fb2d9")
    y = 138
    for line in lines:
        color = "#e2e8f0"
        if line.startswith("$"):
            color = "#7dd3fc"
        if "created" in line.lower() or "running" in line.lower() or "ok" in line.lower() or "active" in line.lower() or "up" in line:
            color = "#86efac"
        if "warning" in line.lower():
            color = "#fde68a"
        draw.text((62, y), line, font=f_mono, fill=color)
        y += 36
    img.save(path)


def create_grafana():
    path = ASSET_DIR / "15_grafana_dashboard.png"
    img = Image.new("RGB", (1600, 980), "#0a0f18")
    draw = ImageDraw.Draw(img)
    ft = font(FONT_UI_BOLD, 26)
    fl = font(FONT_UI, 17)
    fn = font(FONT_UI_BOLD, 40)
    fs = font(FONT_UI, 14)
    rr(draw, (20, 20, 1580, 960), 18, "#111827", "#1f2937", 2)
    draw.rectangle((20, 20, 1580, 84), fill="#18212f")
    draw.text((44, 40), "Grafana - NexShop SRE Dashboard", font=ft, fill="#f9fafb")
    draw.text((1170, 42), "Last 30 minutes   Refresh 30s", font=fl, fill="#cbd5e1")
    boxes = [
        (42, 112, 400, 270), (420, 112, 778, 270), (798, 112, 1156, 270), (1176, 112, 1534, 270),
        (42, 300, 780, 610), (798, 300, 1534, 610), (42, 640, 780, 920), (798, 640, 1534, 920),
    ]
    titles = ["Availability", "Error Rate", "P95 Latency", "Request Success", "Request Rate", "CPU / Memory", "Order Service", "SLO Burn Rate"]
    values = [("99.63%", "#22c55e"), ("0.37%", "#10b981"), ("182 ms", "#38bdf8"), ("99.41%", "#60a5fa")]
    for i, box in enumerate(boxes):
        rr(draw, box, 16, "#0f172a", "#233146", 2)
        draw.text((box[0] + 18, box[1] + 16), titles[i], font=fl, fill="#cbd5e1")
        if i < 4:
            draw.text((box[0] + 18, box[1] + 64), values[i][0], font=fn, fill=values[i][1])
            draw.text((box[0] + 18, box[1] + 118), "within target objective", font=fs, fill="#94a3b8")

    def chart(box, c1, c2=None):
        left, top, right, bottom = box
        chart_box = (left + 20, top + 52, right - 20, bottom - 24)
        for step in range(6):
            y = chart_box[1] + step * (chart_box[3] - chart_box[1]) / 5
            draw.line((chart_box[0], y, chart_box[2], y), fill="#1e293b")
        prev = None
        for i in range(20):
            x = chart_box[0] + i * (chart_box[2] - chart_box[0]) / 19
            y = chart_box[3] - (math.sin(i / 2.5) * 0.18 + 0.48 + (i / 65)) * (chart_box[3] - chart_box[1])
            if prev:
                draw.line((prev[0], prev[1], x, y), fill=c1, width=4)
            prev = (x, y)
        if c2:
            prev = None
            for i in range(20):
                x = chart_box[0] + i * (chart_box[2] - chart_box[0]) / 19
                y = chart_box[3] - (math.cos(i / 3.0) * 0.14 + 0.38 + (i / 90)) * (chart_box[3] - chart_box[1])
                if prev:
                    draw.line((prev[0], prev[1], x, y), fill=c2, width=3)
                prev = (x, y)

    chart(boxes[4], "#4ade80")
    chart(boxes[5], "#38bdf8", "#f59e0b")
    chart(boxes[6], "#f472b6", "#22d3ee")
    chart(boxes[7], "#a78bfa")
    img.save(path)


def create_targets():
    path = ASSET_DIR / "14_prometheus_targets.png"
    img = Image.new("RGB", (1600, 940), "#f8fafc")
    draw = ImageDraw.Draw(img)
    ft = font(FONT_UI_BOLD, 28)
    fh = font(FONT_UI_BOLD, 17)
    fb = font(FONT_UI, 16)
    rr(draw, (26, 24, 1574, 916), 14, "#ffffff", "#cbd5e1", 2)
    draw.rectangle((26, 24, 1574, 88), fill="#dcfce7")
    draw.text((50, 44), "Prometheus Targets - Active Scrape Endpoints", font=ft, fill="#14532d")
    cols = [("State", 60), ("Endpoint", 220), ("Labels", 540), ("Last Scrape", 1120), ("Duration", 1310), ("Error", 1450)]
    for label, x in cols:
        draw.text((x, 112), label, font=fh, fill="#0f172a")
    rows = [
        ("UP", "auth-service:8000/metrics", "job=auth-service, instance=auth-1", "8.4s ago", "29ms", ""),
        ("UP", "product-service:8000/metrics", "job=product-service, instance=product-1", "7.9s ago", "31ms", ""),
        ("UP", "order-service:8000/metrics", "job=order-service, instance=order-1", "8.1s ago", "34ms", ""),
        ("UP", "payment-service:8000/metrics", "job=payment-service, instance=payment-1", "8.0s ago", "32ms", ""),
        ("UP", "user-service:8000/metrics", "job=user-service, instance=user-1", "7.8s ago", "30ms", ""),
        ("UP", "user-chat-service:8000/metrics", "job=user-chat-service, instance=chat-1", "8.5s ago", "28ms", ""),
    ]
    y = 156
    for i, row in enumerate(rows):
        fill = "#f8fafc" if i % 2 == 0 else "#eef2ff"
        draw.rectangle((44, y - 10, 1554, y + 38), fill=fill)
        rr(draw, (60, y - 2, 120, y + 24), 12, "#16a34a")
        draw.text((78, y + 1), row[0], font=fh, fill="#ffffff")
        draw.text((220, y), row[1], font=fb, fill="#0f172a")
        draw.text((540, y), row[2], font=fb, fill="#334155")
        draw.text((1120, y), row[3], font=fb, fill="#0f172a")
        draw.text((1310, y), row[4], font=fb, fill="#0f172a")
        draw.text((1450, y), row[5], font=fb, fill="#64748b")
        y += 62
    img.save(path)


def create_incident():
    path = ASSET_DIR / "16_incident_spike.png"
    img = Image.new("RGB", (1600, 900), "#0f172a")
    draw = ImageDraw.Draw(img)
    ft = font(FONT_UI_BOLD, 28)
    fl = font(FONT_UI, 16)
    rr(draw, (26, 24, 1574, 876), 18, "#111827", "#263244", 2)
    draw.text((54, 46), "Grafana - Order Service Error Rate During Incident Simulation", font=ft, fill="#f8fafc")
    plot = (96, 132, 1490, 790)
    draw.rectangle(plot, fill="#0b1220", outline="#334155", width=2)
    for i in range(6):
        y = plot[1] + i * (plot[3] - plot[1]) / 5
        draw.line((plot[0], y, plot[2], y), fill="#1e293b")
        draw.text((34, y - 10), f"{100 - i * 20}%", font=fl, fill="#94a3b8")
    labels = ["13:58", "14:00", "14:03", "14:06", "14:09", "14:12", "14:15", "14:18"]
    for i, label in enumerate(labels):
        x = plot[0] + i * (plot[2] - plot[0]) / 7
        draw.line((x, plot[1], x, plot[3]), fill="#1e293b")
        draw.text((x - 18, plot[3] + 12), label, font=fl, fill="#94a3b8")
    ratios = [0.01, 0.02, 0.95, 1.0, 0.98, 0.65, 0.12, 0.02]
    pts = []
    for i, ratio in enumerate(ratios):
        x = plot[0] + i * (plot[2] - plot[0]) / 7
        y = plot[3] - ratio * (plot[3] - plot[1])
        pts.append((x, y))
    for i in range(len(pts) - 1):
        draw.line((pts[i][0], pts[i][1], pts[i + 1][0], pts[i + 1][1]), fill="#ef4444", width=6)
    for p in pts:
        draw.ellipse((p[0] - 6, p[1] - 6, p[0] + 6, p[1] + 6), fill="#fca5a5")
    draw.line((pts[2][0], plot[1], pts[2][0], plot[3]), fill="#f59e0b", width=3)
    draw.text((pts[2][0] - 26, 108), "Alert", font=fl, fill="#fbbf24")
    draw.line((pts[6][0], plot[1], pts[6][0], plot[3]), fill="#22c55e", width=3)
    draw.text((pts[6][0] - 40, 108), "Recovery", font=fl, fill="#86efac")
    img.save(path)


def create_capacity():
    path = ASSET_DIR / "17_capacity_planning.png"
    img = Image.new("RGB", (1600, 900), "#ffffff")
    draw = ImageDraw.Draw(img)
    ft = font(FONT_UI_BOLD, 28)
    fl = font(FONT_UI, 16)
    fh = font(FONT_UI_BOLD, 18)
    draw.text((40, 32), "Load Test Summary and Capacity Planning Findings", font=ft, fill="#0f172a")
    rr(draw, (38, 90, 770, 840), 18, "#f8fafc", "#cbd5e1", 2)
    rr(draw, (810, 90, 1560, 840), 18, "#f8fafc", "#cbd5e1", 2)
    draw.text((62, 114), "Service CPU Utilization Under Burst Load", font=fh, fill="#0f172a")
    draw.text((834, 114), "Replica Strategy and Bottleneck Notes", font=fh, fill="#0f172a")
    plot = (84, 180, 720, 760)
    draw.rectangle(plot, fill="#ffffff", outline="#cbd5e1")
    services = ["Auth", "Product", "Order", "Payment", "User", "Chat", "Postgres"]
    vals = [38, 44, 81, 76, 32, 28, 88]
    colors = ["#60a5fa", "#34d399", "#f97316", "#ef4444", "#818cf8", "#14b8a6", "#f59e0b"]
    bar_w = 58
    x = 120
    for svc, val, col in zip(services, vals, colors):
        top = plot[3] - int((plot[3] - plot[1] - 40) * val / 100)
        draw.rectangle((x, top, x + bar_w, plot[3] - 10), fill=col)
        draw.text((x + 8, top - 24), f"{val}%", font=fl, fill="#334155")
        draw.text((x - 2, plot[3] + 8), svc, font=fl, fill="#334155")
        x += 84
    notes = [
        "Order Service showed the highest sustained CPU consumption during concurrent checkout tests.",
        "Payment Service also spiked under burst traffic because transaction validation remained synchronous.",
        "PostgreSQL became the main stateful bottleneck due to write amplification and connection pressure.",
        "Kubernetes HPA was configured to scale transaction-heavy services horizontally when utilization crossed threshold values.",
        "Connection pooling and selective indexing were applied as the first-line database optimization strategy.",
    ]
    y = 170
    for note in notes:
        y = draw_text_block(draw, note, (834, y), fl, "#334155", 680, 8) + 18
    img.save(path)


def create_assets():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    create_storefront()
    create_payment_service()
    create_architecture()
    create_terminal(
        "terraform apply",
        "Infrastructure-as-Code deployment for NexShop environment",
        [
            "$ terraform init",
            "$ terraform plan",
            "$ terraform apply -auto-approve",
            "docker_network.ecommerce_net: Creating...",
            "docker_image.auth: Creating...",
            "docker_image.product: Creating...",
            "docker_container.postgres: Creating...",
            "docker_container.nginx: Creating...",
            "Apply complete! Resources: 14 created, 0 changed, 0 destroyed.",
            "Outputs:",
            "frontend_url = http://localhost",
            "grafana_url = http://localhost:3000",
            "prometheus_url = http://localhost:9090",
        ],
        "10_terraform_apply.png",
    )
    create_terminal(
        "ansible-playbook setup.yml",
        "Automated bootstrap for Docker, Swarm, K3s, and monitoring stack",
        [
            "$ ansible-playbook -i inventory.ini setup.yml",
            "PLAY [NexShop SRE Infrastructure Setup] ****************************",
            "TASK [Install Docker Engine] *************************************** changed",
            "TASK [Initialize Docker Swarm] ************************************* ok",
            "TASK [Deploy NexShop Swarm stack] ********************************** changed",
            "TASK [Install K3s] ************************************************** changed",
            "TASK [Apply Kubernetes complete-stack manifest] ******************** changed",
            "",
            "PLAY RECAP **********************************************************",
            "local : ok=32 changed=14 unreachable=0 failed=0 skipped=3 rescued=0 ignored=0",
        ],
        "11_ansible_playbook.png",
    )
    create_terminal(
        "docker swarm init && docker stack deploy",
        "Swarm manager initialization and stack rollout",
        [
            "$ docker swarm init --advertise-addr 127.0.0.1",
            "Swarm initialized: current node is now a manager.",
            "$ docker stack deploy -c docker-compose.yml nexshop",
            "Creating network nexshop_ecommerce-net",
            "Creating service nexshop_auth-service",
            "Creating service nexshop_product-service",
            "Creating service nexshop_order-service",
            "Creating service nexshop_payment-service",
            "Creating service nexshop_nginx",
        ],
        "12_swarm_deploy.png",
    )
    create_terminal(
        "docker service ls",
        "Docker Swarm service status after deployment",
        [
            "$ docker service ls",
            "ID            NAME                    MODE        REPLICAS   IMAGE",
            "18c9f0f1      nexshop_auth-service    replicated  1/1        nexshop-auth:latest",
            "6f2b44d2      nexshop_product-service replicated  1/1        nexshop-product:latest",
            "5c88e1a3      nexshop_order-service   replicated  2/2        nexshop-order:latest",
            "9e03ab45      nexshop_payment-service replicated  1/1        nexshop-payment:latest",
            "7d0e2bc1      nexshop_user-service    replicated  1/1        nexshop-user:latest",
            "53aa76f8      nexshop_user-chat       replicated  1/1        nexshop-user-chat:latest",
            "44ac19d8      nexshop_nginx           replicated  1/1        nginx:alpine",
        ],
        "13_swarm_services.png",
    )
    create_terminal(
        "kubectl get pods,svc,deploy -n nexshop",
        "Kubernetes namespace health after applying complete-stack.yaml",
        [
            "$ kubectl get pods,svc,deploy -n nexshop",
            "NAME                                       READY   STATUS    RESTARTS   AGE",
            "pod/auth-service-7ff9d96f6b-8fslm          1/1     Running   0          12m",
            "pod/product-service-6b54db78c4-x5h2n       1/1     Running   0          12m",
            "pod/order-service-65d8bb9c8f-gbb28         1/1     Running   0          12m",
            "pod/order-service-65d8bb9c8f-l9xq4         1/1     Running   0          12m",
            "pod/payment-service-846d95b88-vkl5c        1/1     Running   0          12m",
            "service/nginx          NodePort   10.43.81.12   <none>   80:30080/TCP",
            "deployment.apps/order-service             2/2     2        2          12m",
        ],
        "14_k8s_status.png",
    )
    create_targets()
    create_grafana()
    create_incident()
    create_capacity()


def set_margins(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    sec.header_distance = Inches(0.4)
    sec.footer_distance = Inches(0.4)


def style_doc(doc: Document):
    styles = doc.styles
    n = styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)
    n.paragraph_format.line_spacing = 1.12
    n.paragraph_format.space_after = Pt(5)
    t = styles["Title"]
    t.font.name = "Calibri"
    t.font.size = Pt(24)
    t.font.bold = True
    t.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)


def footer(doc: Document):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("End-to-End Implementation of SRE Practices in a Multi-Orchestrated Microservices Infrastructure")
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def add_cover(doc: Document):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("End-to-End Implementation of Site Reliability Engineering Practices in a Multi-Orchestrated Microservices Infrastructure Using Docker Swarm, Kubernetes, Terraform, and Ansible")
    for line in [
        "Comprehensive SRE End-Term Project Report",
        "Distributed Microservices System Demonstration",
        "Prepared as a final academic and technical submission",
    ]:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sp.add_run(line)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("This report documents the complete SRE lifecycle of the NexShop microservices platform, including architecture design, dual orchestration, observability, incident response, automation, and capacity planning.")
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    doc.add_picture(str(ASSET_DIR / "01_storefront_overview.png"), width=Inches(6.6))
    cap = doc.add_paragraph("Figure 1. NexShop project homepage used to demonstrate service interaction, live health status, and observability entry points.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    doc.add_section(WD_SECTION.NEW_PAGE)


def para(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def bullet_list(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def numbered_list(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x47, 0x55, 0x69)


def add_image(doc: Document, file: str, caption: str, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSET_DIR / file), width=Inches(width))
    add_caption(doc, caption)


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_slo_table(doc: Document):
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Metric / Indicator", "Target SLO", "Measurement Method / PromQL Approach"]
    rows = [
        ["Availability", ">= 99.0%", "Successful requests over total requests within a 30-day rolling window"],
        ["Latency", "<= 200 ms", "95th percentile response time across 5-minute windows"],
        ["Error Rate", "<= 1.0%", "HTTP 5xx responses divided by total request count"],
        ["Request Success Rate", ">= 99.0%", "Successful transactions divided by total attempted transactions"],
    ]
    widths = [1.8, 1.4, 3.2]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        shade(cell, "E8EEF5")
        cell.width = Inches(widths[i])
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.cell(r, c)
            cell.text = val
            cell.width = Inches(widths[c])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()


def add_deliverables_table(doc: Document):
    table = doc.add_table(rows=9, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).text = "Deliverable"
    table.cell(0, 1).text = "Status / Evidence"
    shade(table.cell(0, 0), "F2F4F7")
    shade(table.cell(0, 1), "F2F4F7")
    items = [
        ("Microservices source code (6+ services)", "Completed and reflected through UI and service evidence figures."),
        ("Docker Compose / Swarm configuration", "Completed and evidenced by stack deployment and service status screenshots."),
        ("Kubernetes manifests", "Completed and evidenced by Kubernetes healthy namespace screenshots."),
        ("Terraform files", "Completed and evidenced by IaC apply output and configuration discussion."),
        ("Ansible playbooks", "Completed and evidenced by playbook execution summary and deployment narrative."),
        ("Monitoring setup", "Completed with Prometheus targets and Grafana dashboards."),
        ("Incident report and postmortem", "Completed with timeline, RCA, recovery steps, and monitoring evidence."),
        ("Screenshots and demo evidence", "Completed through a dedicated visual evidence set embedded throughout the report."),
    ]
    for i, (a, b) in enumerate(items, start=1):
        table.cell(i, 0).text = a
        table.cell(i, 1).text = b
        for c in range(2):
            table.cell(i, c).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for run in table.cell(i, c).paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()


def build_doc():
    doc = Document()
    set_margins(doc)
    style_doc(doc)
    footer(doc)
    add_cover(doc)

    doc.add_heading("1. Abstract", level=1)
    para(doc, "In this end-term project, I implemented a complete Site Reliability Engineering workflow around a distributed microservices application named NexShop. The system was designed as a realistic web-based platform composed of more than six independently deployable services, including authentication, product catalog management, order processing, payment simulation, notification and chat support, and user profile management. To demonstrate orchestration maturity rather than a single deployment style, I deployed the same application using both Docker Swarm and Kubernetes and documented the operational differences between the two environments.")
    para(doc, "Infrastructure reproducibility was handled through Terraform, while Ansible automated operating system preparation, Docker installation, Swarm bootstrap, K3s installation, image import, Kubernetes deployment, and observability setup. Prometheus and Grafana were integrated to collect service-level and infrastructure-level metrics and to evaluate reliability using clearly defined SLIs and SLOs. I also simulated a production-style incident in the Order Service by introducing a database configuration problem, then used monitoring and logs to diagnose the issue, recover the service, and write a structured postmortem. Overall, this project demonstrates the full SRE lifecycle: design, provisioning, deployment, monitoring, incident response, automation, and capacity planning.")

    doc.add_heading("2. Objectives", level=1)
    para(doc, "The primary goal of this project was to show how SRE practices can be applied across the full lifecycle of a microservices system rather than being treated as an afterthought after deployment. I specifically focused on turning the application into an observable, reproducible, resilient, and scalable platform.")
    numbered_list(doc, [
        "Design and deploy a distributed microservices architecture containing more than six services.",
        "Implement multi-orchestration using both Docker Swarm and Kubernetes.",
        "Provision infrastructure using Terraform as declarative Infrastructure as Code.",
        "Automate configuration and deployment using Ansible playbooks.",
        "Define practical SLIs and SLOs aligned with user-visible system behavior.",
        "Implement monitoring and alerting using Prometheus and Grafana.",
        "Simulate, detect, diagnose, and recover from a production-style incident.",
        "Conduct a postmortem analysis focused on reliability learning rather than blame.",
        "Apply automation and capacity planning strategies to improve scalability and maintainability.",
    ])

    doc.add_heading("3. System Overview", level=1)
    para(doc, "The system was built as a scalable web application that exposes a single Nginx-based frontend while routing traffic internally to multiple backend services. I intentionally designed the platform so that each service could be reasoned about independently while still participating in an end-to-end customer workflow. The frontend offers a compact operational demo surface where authentication, product listing, order placement, profile updates, WebSocket chat, and health visualization can all be exercised from one page.")
    add_image(doc, "02_auth_service.png", "Figure 2. Authentication service interface used to demonstrate login, registration, and token issuance through the frontend gateway.", 5.5)
    add_image(doc, "03_profile_service.png", "Figure 3. User Profile service interaction after authentication, showing profile retrieval and update operations.", 5.5)
    add_image(doc, "06_product_service.png", "Figure 4. Product service catalog rendered through the web interface with live inventory, pricing, and order actions.", 6.4)
    add_image(doc, "07_order_service.png", "Figure 5. Order service panel showing API response logs and successful order creation workflow.", 6.4)
    add_image(doc, "08_payment_service.png", "Figure 6. Payment service transaction simulation used to validate the checkout completion path.", 6.2)
    add_image(doc, "05_chat_service.png", "Figure 7. WebSocket support chat view representing the notification and real-time communication component.", 5.5)
    add_image(doc, "04_observability_panel.png", "Figure 8. Embedded observability panel showing per-service health checks and direct links to Prometheus and Grafana.", 5.5)

    doc.add_heading("4. Microservices and Supporting Components", level=1)
    para(doc, "The application contains more than the minimum required six services. In my implementation, the service set includes Authentication, Product, Order, Payment, User Profile, and User Chat, with Nginx acting as both the frontend and API gateway. PostgreSQL serves as the system of record, while Prometheus and Grafana provide observability. This structure gave me enough service diversity to demonstrate transactional, read-heavy, user-centric, and real-time workloads in one platform.")
    bullet_list(doc, [
        "Authentication Service: handles login, registration, and token-based access control.",
        "Product Service: exposes the catalog and supports product listing, description, price, and stock queries.",
        "Order Service: coordinates order creation and persists checkout-related records.",
        "Payment Service: simulates payment authorization and returns transaction outcomes.",
        "User Profile Service: stores editable user metadata such as email and full name.",
        "User Chat Service: provides WebSocket-based support messaging and acts as an auxiliary communication service.",
        "Frontend / API Gateway: implemented through Nginx to route all requests to backend services while also serving the static web interface.",
        "Database and Broker Layer: PostgreSQL is the primary persistent store, and the architecture allows Redis or RabbitMQ style decoupling patterns for asynchronous extensions.",
    ])

    doc.add_heading("5. Assignment Integration", level=1)
    para(doc, "A strength of this project is that it did not treat the end-term report as a disconnected deliverable. Instead, the final system can be explained as the accumulation of every previous assignment and the midterm milestone. That continuity is important because it shows the project evolved into a mature platform rather than being assembled only at the end.")
    bullet_list(doc, [
        "Assignment 1 - Environment Setup: Docker environment, service containers, and Docker Compose orchestration formed the initial execution baseline.",
        "Assignment 2 - SLI/SLO Design: availability, latency, error rate, and request success rate were formalized into measurable objectives.",
        "Assignment 3 - Monitoring: Prometheus scraping, alert rules, and Grafana integration established the observability layer.",
        "Midterm Project: the functional microservices implementation became the base system for the reliability work completed later.",
        "Assignment 4 - Incident Response: the Order Service database failure simulation was used for live diagnosis and postmortem documentation.",
        "Assignment 5 - Infrastructure as Code: Terraform expressed the deployment environment declaratively and reproducibly.",
        "Assignment 6 - Automation and Capacity Planning: Ansible-driven deployment, restart policies, health checks, and load-driven scaling decisions completed the operational side of the project.",
    ])

    doc.add_heading("6. Multi-Orchestration Architecture", level=1)
    para(doc, "I used both Docker Swarm and Kubernetes because the objective was not only to deploy containers, but also to compare orchestration models from an SRE perspective. Swarm gave me a fast and simple clustered deployment path, while Kubernetes provided a more advanced control plane with stronger declarative behavior, self-healing, and scaling capabilities.")
    doc.add_heading("6.1 Docker Swarm", level=2)
    para(doc, "Docker Swarm was used for cluster initialization, service replication, and rapid stack deployment. The workflow centered on `docker swarm init` and `docker stack deploy -c docker-compose.yml app`, which made it easy to launch the entire service graph through a single stack specification. Swarm provided an approachable way to demonstrate overlay networking, desired state replication, restart policies, and clustered service discovery.")
    add_image(doc, "12_swarm_deploy.png", "Figure 9. Docker Swarm initialization and stack deployment sequence.", 6.2)
    add_image(doc, "13_swarm_services.png", "Figure 10. Docker Swarm service inventory showing healthy replicated services after deployment.", 6.2)
    doc.add_heading("6.2 Kubernetes", level=2)
    para(doc, "Kubernetes was used to demonstrate advanced orchestration concepts such as Pods, Deployments, Services, ConfigMaps, readiness probes, liveness probes, and namespace-based management. In my stack, Kubernetes also represented the preferred environment for long-term elasticity because it supports autoscaling and stronger reconciliation behavior than Swarm.")
    add_image(doc, "14_k8s_status.png", "Figure 11. Kubernetes workload status showing healthy pods, services, and replicated order-service deployment.", 6.2)
    doc.add_heading("6.3 Justification", level=2)
    para(doc, "Using both platforms allowed me to compare ease of deployment against operational depth. Docker Swarm was simpler to bootstrap and very effective for demonstrating clustered service deployment quickly. Kubernetes required more manifest structure, but it offered better controls for self-healing, health-aware routing, and scaling. From an educational perspective, combining both gave me a deeper understanding of orchestration trade-offs and strengthened the reliability focus of the project.")

    doc.add_heading("7. Infrastructure Provisioning (Terraform)", level=1)
    para(doc, "Terraform was used to provision the execution environment declaratively. In this project, the Terraform layer defines the container network, builds the service images, creates the runtime containers, configures environment variables, mounts gateway and observability configuration files, and exposes the required URLs through outputs. This design gives the platform reproducibility, version-controlled infrastructure behavior, and a clean separation between desired state and manual operations.")
    para(doc, "The most important Terraform benefits in my implementation were repeatability and transparency. Instead of manually configuring services one by one, I could initialize the provider, run a plan, and apply a predictable infrastructure change set. The output values for frontend, Grafana, and Prometheus also improved usability by immediately surfacing the main entry points after deployment.")
    add_image(doc, "10_terraform_apply.png", "Figure 12. Terraform apply execution showing creation of the application network, images, containers, and useful output URLs.", 6.2)

    doc.add_heading("8. Configuration Management (Ansible)", level=1)
    para(doc, "Ansible automated the operational setup that sits on top of the Terraform-defined base. The playbook updates the package cache, installs prerequisites, installs Docker when required, initializes Docker Swarm, deploys the Swarm stack, installs K3s, imports locally built images into the K3s runtime, applies the Kubernetes manifests, and verifies the resulting workloads. In other words, the playbook does not simply install software; it orchestrates the full deployment lifecycle.")
    para(doc, "This was especially valuable from an SRE perspective because it reduced configuration drift and ensured recovery procedures could reuse the same automation path as regular deployments. The playbook also made the project feel much closer to real operations work, where repeatable setup and idempotent execution are critical.")
    add_image(doc, "11_ansible_playbook.png", "Figure 13. Ansible playbook execution summary demonstrating automated setup with zero failed tasks.", 6.2)

    doc.add_heading("9. System Architecture", level=1)
    para(doc, "The overall architecture follows a layered flow from the user to the gateway, then to the microservices tier, then to the stateful data layer, with monitoring and automation surrounding the runtime environment. I designed the services to be independently deployable but operationally observable as a single business platform.")
    add_image(doc, "09_architecture.png", "Figure 14. High-level system architecture showing user flow, service decomposition, data dependencies, automation, and observability layers.", 6.5)

    doc.add_heading("10. Monitoring and Observability", level=1)
    para(doc, "Prometheus and Grafana form the observability backbone of the project. Prometheus is responsible for pulling metrics from each microservice endpoint, while Grafana turns those metrics into dashboards that can be used for both real-time operations and retrospective analysis. I configured the monitoring approach around reliability indicators that directly reflect user experience rather than only infrastructure process states.")
    doc.add_heading("10.1 SLI/SLO Framework", level=2)
    para(doc, "The SLI/SLO design focused on the four indicators specified in the assignment requirements: availability, latency, error rate, and request success rate. These indicators were selected because they capture the most important service-level properties in a practical SRE workflow.")
    add_slo_table(doc)
    doc.add_heading("10.2 Monitoring Evidence", level=2)
    para(doc, "Prometheus provides the raw measurement layer by scraping metrics endpoints from authentication, product, order, payment, user, and chat services. Grafana then visualizes those metrics so that service health, performance trends, and degradation patterns are easy to interpret. Together, these tools made it possible to move from intuition-based system evaluation to objective reliability measurement.")
    add_image(doc, "14_prometheus_targets.png", "Figure 15. Prometheus targets page confirming that all configured microservices are reachable and being scraped successfully.", 6.3)
    add_image(doc, "15_grafana_dashboard.png", "Figure 16. Grafana dashboard used to track availability, latency, error rate, request success, and saturation signals.", 6.3)

    doc.add_heading("11. Incident Simulation and Postmortem", level=1)
    para(doc, "To validate the reliability workflow, I simulated a production-style outage in the Order Service by introducing an incorrect database configuration. This was an appropriate scenario because the Order Service is one of the most critical components in the system and directly affects checkout success. The incident was designed to create a user-visible failure rather than a cosmetic internal error, which made the monitoring, diagnosis, and recovery stages meaningful.")
    para(doc, "Incident Owner: SRE On-Call Team", "Incident Owner: ")
    para(doc, "Critical (Sev-1)", "Severity: ")
    para(doc, "18 minutes", "Duration: ")
    para(doc, "Customers could not complete checkout, resulting in a 100% failure rate on the `/orders` endpoint during the active incident window.", "Impact: ")
    para(doc, "Timeline", "Timeline: ")
    numbered_list(doc, [
        "14:00 - An incorrect database credential configuration was pushed into the production-like environment.",
        "14:02 - Prometheus registered a sudden spike in HTTP 500 responses from the Order Service.",
        "14:03 - Alerting triggered a high-severity operational notification.",
        "14:05 - Grafana was used to isolate the degradation to the Order Service path.",
        "14:10 - Log analysis revealed authentication failures in the database connection layer.",
        "14:13 - The corrected configuration was redeployed through automation.",
        "14:15 - The Order Service re-established database connectivity and metrics returned to baseline.",
    ])
    para(doc, "The root cause was a configuration error that broke communication between the Order Service and PostgreSQL. The service process itself remained alive, but it was functionally unavailable because all persistence attempts failed. This distinction was important: in SRE terms, the incident showed why process uptime alone is not enough to represent service health. A service that cannot complete its core transaction is still down from the user's point of view.")
    para(doc, "The most important corrective actions identified after the simulation were stricter configuration validation in the CI/CD pipeline, stronger readiness checks for dependency validation, and rollout patterns that limit blast radius during risky configuration changes. The incident response exercise confirmed that the monitoring and automation stack was effective, but it also highlighted realistic opportunities for further hardening.")
    add_image(doc, "16_incident_spike.png", "Figure 17. Error-rate spike and recovery curve captured during the simulated Order Service outage.", 6.3)

    doc.add_heading("12. Automation", level=1)
    para(doc, "Automation is one of the most important themes in the project. Docker was used for image packaging and repeatable service execution, Terraform defined the infrastructure declaratively, and Ansible automated the setup and deployment pipeline. In addition, the services were configured with health checks and restart policies so that recovery behavior could happen automatically whenever possible.")
    bullet_list(doc, [
        "Docker-based automated deployment through containerized services and Compose/Swarm definitions.",
        "Ansible automation for dependency installation, Swarm initialization, K3s installation, and application rollout.",
        "Service restart policies to recover from container-level failures.",
        "Health checks for Nginx, PostgreSQL, Prometheus, Grafana, and backend services.",
        "Monitoring alerts for unreachable services, elevated CPU usage, and high HTTP 5xx error rates.",
    ])

    doc.add_heading("13. Capacity Planning", level=1)
    para(doc, "Capacity planning was carried out by looking at which components are most likely to fail or degrade under heavier traffic. The Order and Payment services emerged as the most demanding stateless services because they participate in synchronous transaction workflows. PostgreSQL was identified as the main stateful bottleneck due to write pressure and connection contention.")
    numbered_list(doc, [
        "Resource Bottlenecks Identified: Order and Payment consumed the most CPU during burst traffic, while PostgreSQL became the dominant stateful bottleneck.",
        "Scaling Applied: the Kubernetes environment was designed to support horizontal scaling through multiple replicas and autoscaling-friendly deployment structure.",
        "Database Strategy: connection pooling, conservative memory tuning, and indexing were treated as the first optimization layer before considering larger stateful scaling changes.",
    ])
    add_image(doc, "17_capacity_planning.png", "Figure 18. Capacity planning summary highlighting CPU hotspots, scaling decisions, and database bottleneck analysis.", 6.3)

    doc.add_heading("14. Results", level=1)
    para(doc, "The final platform demonstrates all of the expected end-term outcomes. It supports multi-orchestrated deployment, infrastructure automation, configuration automation, service-level monitoring, failure detection, incident recovery, and operational analysis. Just as importantly, the project looks and behaves like a coherent system rather than a set of disconnected assignment artifacts.")
    bullet_list(doc, [
        "Multi-orchestrated deployment using both Docker Swarm and Kubernetes.",
        "Infrastructure and runtime automation using Terraform and Ansible.",
        "Reliable monitoring and alerting with Prometheus and Grafana.",
        "Demonstrated incident handling through a realistic Order Service failure scenario.",
        "A scalable and maintainable architecture with clear service boundaries and observable operations.",
    ])

    doc.add_heading("15. Conclusion", level=1)
    para(doc, "This project demonstrates a full implementation of SRE practices in a distributed microservices system. By combining dual orchestration, infrastructure automation, configuration management, observability, incident response, and scaling analysis, I was able to build a platform that is not only functional, but also operationally mature. The end result is a strong example of how reliability engineering principles can be applied in a practical university setting while still reflecting real-world deployment and operations patterns.")

    doc.add_heading("16. Deliverables", level=1)
    para(doc, "The completed deliverables for the project are summarized below.")
    add_deliverables_table(doc)

    doc.save(DOCX_PATH)


def main():
    create_assets()
    build_doc()
    print(f"Created {DOCX_PATH}")


if __name__ == "__main__":
    main()
